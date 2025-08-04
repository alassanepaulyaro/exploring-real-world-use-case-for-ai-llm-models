# Professional Meeting Documentation with Advanced AI Processing
import streamlit as st
import requests
import speech_recognition as sr
import tempfile
import os
from datetime import datetime
import json
import re
from langdetect import detect
import base64
from io import BytesIO
import time
from pydub import AudioSegment
from pydub.utils import make_chunks
from fpdf import FPDF
from docx import Document
from docx.shared import Inches

# Configuration
OLLAMA_URL = "http://localhost:11434/api/generate"
OLLAMA_API_URL = "http://localhost:11434/api/tags"
DEFAULT_MODEL = "mistral-nemo:latest"

class MeetingMinutesGenerator:
    def __init__(self):
        self.supported_languages = {
            'en': 'English',
            'fr': 'French',
            'es': 'Spanish',
            'de': 'German',
            'it': 'Italian',
            'pt': 'Portuguese',
            'nl': 'Dutch',
            'ru': 'Russian',
            'ja': 'Japanese',
            'ko': 'Korean',
            'zh': 'Chinese'
        }
        
        self.transcription_languages = {
            'en-US': 'English (US)',
            'en-GB': 'English (UK)',
            'fr-FR': 'French (France)',
            'es-ES': 'Spanish (Spain)',
            'es-MX': 'Spanish (Mexico)',
            'de-DE': 'German (Germany)',
            'it-IT': 'Italian (Italy)',
            'pt-BR': 'Portuguese (Brazil)',
            'pt-PT': 'Portuguese (Portugal)',
            'nl-NL': 'Dutch (Netherlands)',
            'ru-RU': 'Russian (Russia)',
            'ja-JP': 'Japanese (Japan)',
            'ko-KR': 'Korean (South Korea)',
            'zh-CN': 'Chinese (Mandarin, China)',
            'zh-TW': 'Chinese (Traditional, Taiwan)'
        }
        
        self.output_formats = {
            'structured': "Create structured meeting minutes with clear sections, bullet points, and professional formatting.",
            'detailed': "Generate comprehensive meeting minutes with detailed discussions, full context, and extensive notes.",
            'executive': "Create executive-level summary focusing on key decisions, high-level outcomes, and strategic points.",
            'action_focused': "Focus primarily on action items, decisions, and next steps with responsible parties and deadlines.",
            'timeline': "Create a chronological timeline of the meeting with time-stamped key events and discussions."
        }
        
        self.meeting_types = {
            'general': "General business meeting with standard corporate structure",
            'project': "Project meeting focused on deliverables, milestones, and team coordination",
            'strategic': "Strategic planning meeting with decision-making and future planning",
            'standup': "Daily standup or quick status meeting format",
            'board': "Board meeting with formal governance and decision documentation",
            'client': "Client meeting focused on relationship management and service delivery"
        }

    def check_ollama_connection(self):
        try:
            response = requests.get(OLLAMA_API_URL, timeout=10)
            if response.status_code != 200:
                return False, "Ollama service not responding"
            
            models = response.json().get('models', [])
            model_names = [model.get('name', '') for model in models]
            
            if not any(DEFAULT_MODEL in name for name in model_names):
                available_models = ', '.join(model_names) if model_names else "None"
                return False, f"Model {DEFAULT_MODEL} not found. Available models: {available_models}"
            
            return True, "Connection successful"
        except requests.exceptions.Timeout:
            return False, "Timeout connecting to Ollama service"
        except requests.exceptions.ConnectionError:
            return False, "Cannot connect to Ollama service. Is it running on localhost:11434?"
        except Exception as e:
            return False, f"Unexpected error: {str(e)}"

    def detect_language(self, text):
        try:
            detected_lang = detect(text)
            return detected_lang if detected_lang in self.supported_languages else 'en'
        except:
            return 'en'

    def extract_participants(self, transcript):
        participants = set()
        patterns = [
            r'([A-Z][a-z]+):\s',
            r'([A-Z][a-z]+)\s+said',
            r'([A-Z][a-z]+)\s+mentioned',
            r'([A-Z][a-z]+)\s+asked',
            r'([A-Z][a-z]+)\s+noted',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, transcript)
            participants.update(matches)
        
        common_words = {'The', 'This', 'That', 'Then', 'Next', 'First', 'Last', 'After', 'Before'}
        participants = participants - common_words
        
        return list(participants)

    def analyze_transcript(self, transcript):
        words = transcript.split()
        sentences = transcript.split('.')
        participants = self.extract_participants(transcript)
        
        speaking_distribution = {}
        for participant in participants:
            count = transcript.lower().count(participant.lower() + ':')
            speaking_distribution[participant] = count
        
        return {
            'word_count': len(words),
            'sentence_count': len(sentences),
            'participant_count': len(participants),
            'participants': participants,
            'speaking_distribution': speaking_distribution,
            'estimated_duration': f"{len(words) // 150}-{len(words) // 100} minutes"
        }

    def generate_structured_prompt(self, transcript, language, output_format, meeting_type, custom_instructions=""):
        detected_lang = self.detect_language(transcript)
        language_instruction = f"Generate the meeting minutes in {self.supported_languages.get(detected_lang, 'English')}."
        
        format_instruction = self.output_formats.get(output_format, self.output_formats['structured'])
        meeting_context = self.meeting_types.get(meeting_type, self.meeting_types['general'])
        
        participants = self.extract_participants(transcript)
        participant_list = ", ".join(participants) if participants else "Multiple participants"
        
        base_prompt = f"""
Create professional meeting minutes from the following transcript with these specifications:

MEETING CONTEXT:
- Type: {meeting_context}
- Participants detected: {participant_list}
- Language: {language_instruction}

TRANSCRIPT:
{transcript}

FORMATTING REQUIREMENTS:
- {format_instruction}
- Include clear section headers
- Extract and highlight key decisions made
- Identify action items with responsible parties where possible
- Note any deadlines or timelines mentioned
- Summarize main discussion points
- Maintain professional tone and structure

ADDITIONAL INSTRUCTIONS:
{custom_instructions}

STRUCTURE THE OUTPUT WITH THESE SECTIONS:
1. Meeting Overview (Date, Participants, Purpose)
2. Key Discussion Points
3. Decisions Made
4. Action Items
5. Next Steps
6. Follow-up Items

Ensure the minutes are professional, comprehensive, and actionable.
"""
        return base_prompt

    def generate_fallback_minutes(self, transcript, meeting_type):
        participants = self.extract_participants(transcript)
        analysis = self.analyze_transcript(transcript)
        
        return f"""
===========================================
MEETING MINUTES
===========================================

Meeting Date: {datetime.now().strftime("%Y-%m-%d")}
Meeting Type: {meeting_type.title()}
Participants: {', '.join(participants) if participants else 'Multiple participants'}
Duration: {analysis['estimated_duration']} (estimated)

===========================================
OVERVIEW
===========================================

This meeting involved {analysis['participant_count']} participants discussing various topics.
Total words in transcript: {analysis['word_count']}

===========================================
KEY DISCUSSION POINTS
===========================================

[Based on transcript analysis]
• Main topics were discussed among participants
• Various viewpoints and perspectives were shared
• Multiple agenda items were covered during the session

===========================================
PARTICIPANTS CONTRIBUTION
===========================================

{chr(10).join([f"• {name}: {count} speaking instances" for name, count in analysis['speaking_distribution'].items()])}

===========================================
ACTION ITEMS
===========================================

• Review transcript for specific action items
• Follow up on decisions mentioned in the discussion
• Schedule next meeting if required

===========================================
NEXT STEPS
===========================================

• Distribute meeting minutes to all participants
• Track progress on identified action items
• Prepare for follow-up discussions as needed

Note: This is a basic template. For detailed minutes, please ensure AI service is available.
"""

    def generate_meeting_minutes(self, transcript, language, output_format, meeting_type, custom_instructions=""):
        is_connected, connection_message = self.check_ollama_connection()
        
        if not is_connected:
            st.warning(f"AI Service Issue: {connection_message}")
            st.info("Using fallback template generation...")
            return self.generate_fallback_minutes(transcript, meeting_type)
        
        prompt = self.generate_structured_prompt(
            transcript, language, output_format, meeting_type, custom_instructions
        )
        
        payload = {
            "model": DEFAULT_MODEL,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": 0.3,
                "top_p": 0.9,
                "num_predict": 3000
            }
        }
        
        try:
            response = requests.post(OLLAMA_URL, json=payload, timeout=120)
            
            if response.status_code == 200:
                result = response.json().get("response", "")
                if result:
                    return result
                else:
                    st.warning("Empty response from AI. Using fallback template...")
                    return self.generate_fallback_minutes(transcript, meeting_type)
            else:
                st.error(f"AI Service Error: {response.status_code} - {response.text}")
                return self.generate_fallback_minutes(transcript, meeting_type)
        except requests.exceptions.Timeout:
            st.warning("AI generation timed out. Using fallback template...")
            return self.generate_fallback_minutes(transcript, meeting_type)
        except requests.exceptions.RequestException as e:
            st.error(f"Connection Error: {str(e)}")
            return self.generate_fallback_minutes(transcript, meeting_type)

    def clean_text_for_pdf(self, text):
        unicode_replacements = {
            '\u0153': 'oe', '\u0152': 'OE', '\u00e0': 'a', '\u00e1': 'a', '\u00e2': 'a', '\u00e4': 'a',
            '\u00e8': 'e', '\u00e9': 'e', '\u00ea': 'e', '\u00eb': 'e', '\u00ec': 'i', '\u00ed': 'i',
            '\u00ee': 'i', '\u00ef': 'i', '\u00f2': 'o', '\u00f3': 'o', '\u00f4': 'o', '\u00f6': 'o',
            '\u00f9': 'u', '\u00fa': 'u', '\u00fb': 'u', '\u00fc': 'u', '\u00f1': 'n', '\u00e7': 'c',
            '\u2013': '-', '\u2014': '-', '\u2018': "'", '\u2019': "'", '\u201c': '"', '\u201d': '"',
            '\u2022': '•', '\u2026': '...'
        }
        
        cleaned_text = text
        for unicode_char, replacement in unicode_replacements.items():
            cleaned_text = cleaned_text.replace(unicode_char, replacement)
        
        cleaned_text = ''.join(char if ord(char) < 128 else '?' for char in cleaned_text)
        return cleaned_text

    def export_to_pdf(self, content, filename):
        try:
            clean_content = self.clean_text_for_pdf(content)
            
            class PDF(FPDF):
                def header(self):
                    self.set_font('Arial', 'B', 18)
                    self.set_text_color(44, 62, 80)
                    self.cell(0, 15, 'Meeting Minutes', 0, 1, 'C')
                    self.set_font('Arial', '', 10)
                    self.set_text_color(128, 128, 128)
                    self.cell(0, 10, f'Generated on {datetime.now().strftime("%Y-%m-%d at %H:%M")}', 0, 1, 'C')
                    self.set_text_color(0, 0, 0)
                    self.ln(10)
                
                def footer(self):
                    self.set_y(-15)
                    self.set_font('Arial', 'I', 8)
                    self.set_text_color(128, 128, 128)
                    self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')
            
            pdf = PDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            
            lines = clean_content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    pdf.ln(4)
                    continue
                
                if line.startswith('=') and line.endswith('='):
                    continue
                elif line.isupper() and len(line.split()) <= 8:
                    pdf.set_font('Arial', 'B', 14)
                    pdf.set_text_color(44, 62, 80)
                    pdf.cell(0, 10, line, 0, 1, 'L')
                    pdf.set_text_color(0, 0, 0)
                    pdf.ln(3)
                elif line.startswith('•') or line.startswith('-'):
                    pdf.set_font('Arial', '', 10)
                    if len(line) > 85:
                        words = line.split()
                        current_line = ""
                        first_line = True
                        for word in words:
                            if len(current_line + " " + word) <= 85:
                                current_line += (" " + word) if current_line else word
                            else:
                                if current_line:
                                    indent = "" if first_line else "  "
                                    pdf.cell(0, 6, indent + current_line, 0, 1, 'L')
                                    first_line = False
                                current_line = "  " + word if not first_line else word
                        if current_line:
                            indent = "" if first_line else "  "
                            pdf.cell(0, 6, indent + current_line, 0, 1, 'L')
                    else:
                        pdf.cell(0, 6, line, 0, 1, 'L')
                    pdf.ln(1)
                else:
                    pdf.set_font('Arial', '', 11)
                    if len(line) > 90:
                        words = line.split()
                        current_line = ""
                        for word in words:
                            if len(current_line + " " + word) <= 90:
                                current_line += (" " + word) if current_line else word
                            else:
                                if current_line:
                                    pdf.cell(0, 7, current_line, 0, 1, 'L')
                                current_line = word
                        if current_line:
                            pdf.cell(0, 7, current_line, 0, 1, 'L')
                    else:
                        pdf.cell(0, 7, line, 0, 1, 'L')
                    pdf.ln(2)
            
            pdf_output = pdf.output(dest='S')
            return pdf_output.encode('latin-1') if isinstance(pdf_output, str) else pdf_output
        except Exception as e:
            st.error(f"PDF export failed: {str(e)}")
            return None

    def export_to_docx(self, content, filename):
        try:
            doc = Document()
            
            title = doc.add_heading('Meeting Minutes', 0)
            title.alignment = 1
            
            date_para = doc.add_paragraph(f'Generated on {datetime.now().strftime("%Y-%m-%d at %H:%M")}')
            date_para.alignment = 1
            
            doc.add_paragraph()
            
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                
                if line.startswith('=') and line.endswith('='):
                    continue
                elif line.isupper() and len(line.split()) <= 8:
                    doc.add_heading(line, level=1)
                elif line.startswith('•') or line.startswith('-'):
                    para = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                else:
                    doc.add_paragraph(line)
            
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer.getvalue()
        except Exception as e:
            st.error(f"DOCX export failed: {str(e)}")
            return None

    def convert_audio_to_wav(self, audio_file, target_sample_rate=16000):
        try:
            audio_file.seek(0)
            file_extension = audio_file.name.split('.')[-1].lower()
            
            if file_extension == 'wav':
                with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as tmp_file:
                    tmp_file.write(audio_file.read())
                    tmp_file.flush()
                    return tmp_file.name, None
            
            try:
                test_audio = AudioSegment.empty()
                test_audio.export(format="wav")
            except Exception:
                return None, ("FFmpeg not found. Please install FFmpeg:\n"
                            "Windows: Download from https://ffmpeg.org/download.html\n"
                            "Or use: winget install ffmpeg\n"
                            "Linux: sudo apt install ffmpeg\n"
                            "macOS: brew install ffmpeg\n\n"
                            "For now, please use WAV files directly.")
            
            if file_extension == 'mp3':
                audio = AudioSegment.from_mp3(audio_file)
            elif file_extension in ['mp4', 'm4a']:
                audio = AudioSegment.from_file(audio_file, format='mp4')
            elif file_extension == 'flac':
                audio = AudioSegment.from_file(audio_file, format='flac')
            else:
                audio = AudioSegment.from_file(audio_file)
            
            audio = audio.set_channels(1)
            audio = audio.set_frame_rate(target_sample_rate)
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as tmp_wav:
                audio.export(tmp_wav.name, format='wav')
                return tmp_wav.name, None
        except Exception as e:
            return None, f"Audio conversion failed: {str(e)}"

    def transcribe_audio_chunks(self, wav_file_path, language='en-US', chunk_duration=30):
        recognizer = sr.Recognizer()
        
        try:
            audio = AudioSegment.from_wav(wav_file_path)
            chunk_length_ms = chunk_duration * 1000
            chunks = make_chunks(audio, chunk_length_ms)
            
            transcriptions = []
            total_chunks = len(chunks)
            
            for i, chunk in enumerate(chunks):
                with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as chunk_file:
                    chunk.export(chunk_file.name, format='wav')
                    chunk_file_path = chunk_file.name
                
                try:
                    with sr.AudioFile(chunk_file_path) as source:
                        recognizer.adjust_for_ambient_noise(source, duration=0.5)
                        audio_data = recognizer.record(source)
                    
                    chunk_text = recognizer.recognize_google(audio_data, language=language)
                    if chunk_text.strip():
                        transcriptions.append(chunk_text)
                    
                    if 'st' in globals():
                        progress = (i + 1) / total_chunks
                        st.progress(progress, text=f"Processing chunk {i+1}/{total_chunks}")
                except sr.UnknownValueError:
                    continue
                except Exception as e:
                    st.warning(f"Error processing chunk {i+1}: {str(e)}")
                    continue
                finally:
                    os.unlink(chunk_file_path)
            
            full_transcript = ' '.join(transcriptions)
            return full_transcript, "Success"
        except Exception as e:
            return "", f"Chunked transcription failed: {str(e)}"

    def transcribe_audio(self, audio_file, language='en-US'):
        try:
            wav_file_path, conversion_error = self.convert_audio_to_wav(audio_file)
            
            if conversion_error:
                return "", conversion_error
            
            try:
                audio_segment = AudioSegment.from_wav(wav_file_path)
                duration_seconds = len(audio_segment) / 1000
                
                if duration_seconds > 60:
                    st.info(f"📊 Audio duration: {duration_seconds:.1f}s - Using chunked transcription for better accuracy")
                    transcript, error = self.transcribe_audio_chunks(wav_file_path, language)
                else:
                    recognizer = sr.Recognizer()
                    with sr.AudioFile(wav_file_path) as source:
                        recognizer.adjust_for_ambient_noise(source, duration=1)
                        audio_data = recognizer.record(source)
                    transcript = recognizer.recognize_google(audio_data, language=language)
                    error = "Success"
            except Exception as e:
                transcript, error = self.transcribe_audio_chunks(wav_file_path, language)
            
            os.unlink(wav_file_path)
            return transcript, error
        except sr.UnknownValueError:
            return "", "Could not understand audio - please check audio quality and language settings"
        except sr.RequestError as e:
            return "", f"Speech recognition service error: {str(e)}"
        except Exception as e:
            return "", f"Transcription error: {str(e)}"

def main():
    st.set_page_config(
        page_title="AI Meeting Minutes Generator",
        page_icon="📝",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    st.markdown("""
    <style>
    .main-header { text-align: center; padding: 2rem 0; background: linear-gradient(90deg, #667eea 0%, #764ba2 100%); color: white; border-radius: 10px; margin-bottom: 2rem; }
    .metric-card { background: #f8f9fa; padding: 1rem; border-radius: 8px; border-left: 4px solid #667eea; }
    .export-section { background: #f8f9fa; padding: 1.5rem; border-radius: 10px; margin: 1rem 0; }
    .success-banner { background: #d4edda; border: 1px solid #c3e6cb; color: #155724; padding: 1rem; border-radius: 8px; margin: 1rem 0; }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <div class="main-header">
        <h1>📝 AI-Powered Meeting Minutes Generator</h1>
        <p>Professional Meeting Documentation with Advanced AI Processing</p>
    </div>
    """, unsafe_allow_html=True)
    
    generator = MeetingMinutesGenerator()
    
    with st.sidebar:
        st.markdown("### 🔧 System Status")
        
        with st.spinner("Checking AI service..."):
            is_connected, status_message = generator.check_ollama_connection()
        
        if is_connected:
            st.success("✅ AI Service: Connected")
        else:
            st.error(f"❌ AI Service: {status_message}")
            with st.expander("🔧 Troubleshooting"):
                st.markdown("""
                **To fix AI connection issues:**
                1. **Start Ollama service:** `ollama serve`
                2. **Install the model:** `ollama pull mistral-nemo:latest`
                3. **Verify installation:** `ollama list`
                **Note:** The app will use template generation if AI is unavailable.
                """)
        
        st.divider()
        st.markdown("### ⚙️ Configuration")
        
        selected_language = st.selectbox(
            "Output Language",
            options=list(generator.supported_languages.keys()),
            format_func=lambda x: generator.supported_languages[x]
        )
        
        output_format = st.selectbox(
            "Minutes Format",
            options=list(generator.output_formats.keys()),
            format_func=lambda x: x.replace('_', ' ').title(),
            help="Choose the style and detail level for your meeting minutes"
        )
        
        meeting_type = st.selectbox(
            "Meeting Type",
            options=list(generator.meeting_types.keys()),
            format_func=lambda x: x.replace('_', ' ').title(),
            help="Select the type of meeting for optimized formatting"
        )
        
        custom_instructions = st.text_area(
            "Custom Instructions (Optional)",
            placeholder="e.g., Focus on technical decisions, Include specific terminology, etc.",
            height=100
        )

    tab1, tab2, tab3 = st.tabs(["📝 Text Input", "🎤 Audio Upload", "📊 Analysis"])
    
    with tab1:
        st.markdown("### Meeting Transcript Input")
        
        if st.button("📋 Load Sample Transcript", type="secondary"):
            sample_transcript = """John: Good morning everyone. Let's start with the Q4 review.
Sarah: The sales numbers are looking strong. We achieved 15% growth compared to Q3.
Mike: That's excellent. What about the marketing budget for next quarter?
Sarah: We need to increase it by 20% to maintain momentum.
John: I think that's reasonable. Let's approve the budget increase.
Lisa: I'll prepare the detailed budget breakdown by Friday.
Mike: Don't forget we have the client presentation next week.
John: Right, Sarah will handle that. Meeting adjourned."""
            st.session_state.transcribed_text = sample_transcript
            st.rerun()
        
        default_value = st.session_state.get('transcribed_text', '')
        
        transcript = st.text_area(
            "Meeting Transcript *",
            value=default_value,
            placeholder="Paste your meeting transcript here...\n\nExample:\nJohn: We need to discuss the Q4 budget.\nSarah: I agree. Let's review the numbers...",
            height=350,
            key="transcript_input"
        )
        
        if transcript:
            col1, col2, col3 = st.columns(3)
            col1.metric("Characters", len(transcript))
            col2.metric("Words", len(transcript.split()))
            col3.metric("Est. Reading", f"{len(transcript.split()) // 200 + 1} min")
    
    with tab2:
        st.markdown("### Audio Transcription")
        
        st.info("📌 **Audio Requirements:** WAV, MP3, MP4, M4A files. Best results with clear audio and minimal background noise.")
        
        col1, col2 = st.columns(2)
        with col1:
            transcription_language = st.selectbox(
                "Transcription Language",
                options=list(generator.transcription_languages.keys()),
                format_func=lambda x: generator.transcription_languages[x],
                help="Select the language spoken in the audio file"
            )
        
        with col2:
            auto_detect_lang = st.checkbox(
                "Auto-detect language after transcription",
                value=True,
                help="Automatically detect the language of the transcribed text"
            )
        
        audio_file = st.file_uploader(
            "Upload Audio File",
            type=['wav', 'mp3', 'mp4', 'm4a', 'flac'],
            help="Upload your meeting recording for automatic transcription"
        )
        
        if audio_file is not None:
            st.audio(audio_file)
            
            file_size = len(audio_file.getvalue()) / 1024
            st.caption(f"📁 File: {audio_file.name} ({file_size:.1f} KB)")
            
            if st.button("🔄 Transcribe Audio", type="primary"):
                with st.spinner("Transcribing audio... This may take a few minutes."):
                    transcript_result, error_message = generator.transcribe_audio(audio_file, transcription_language)
                
                if transcript_result:
                    st.success("✅ Transcription completed!")
                    
                    if auto_detect_lang:
                        detected_lang = generator.detect_language(transcript_result)
                        detected_lang_name = generator.supported_languages.get(detected_lang, 'Unknown')
                        st.info(f"🌐 Detected language: {detected_lang_name}")
                    
                    st.session_state.transcribed_text = transcript_result
                    
                    st.markdown("### 📝 Review & Edit Transcription")
                    st.info("💡 **Tip:** Review the transcription below and make any necessary corrections before generating minutes.")
                    
                    edited_transcript = st.text_area(
                        "Transcribed Text (Editable)",
                        value=transcript_result,
                        height=300,
                        help="Review and edit the transcription if needed",
                        key="edited_transcript"
                    )
                    
                    if edited_transcript != transcript_result:
                        st.session_state.transcribed_text = edited_transcript
                        st.success("✅ Transcript updated with your edits!")
                    
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Words", len(edited_transcript.split()))
                    col2.metric("Characters", len(edited_transcript))
                    col3.metric("Estimated Reading Time", f"{len(edited_transcript.split()) // 200 + 1} min")
                    
                    if st.button("✅ Use This Transcription", type="secondary"):
                        st.success("✅ Transcription ready! Go to the Text Input tab to see it.")
                        st.rerun()
                    
                else:
                    st.error(f"❌ Transcription failed: {error_message}")
                    with st.expander("💡 Troubleshooting Tips"):
                        st.markdown("""
                        - Ensure clear audio quality with minimal background noise
                        - Try converting your file to WAV format first
                        - Check if the selected language matches the audio
                        - For long files, ensure stable internet connection
                        - Consider splitting very long audio files
                        """)
                    st.info("**Alternative:** You can manually enter the transcript in the 'Text Input' tab.")
    
    with tab3:
        if transcript:
            st.markdown("### Transcript Analysis")
            
            analysis = generator.analyze_transcript(transcript)
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.markdown(f"""
                <div class="metric-card">
                    <h3>{analysis['word_count']}</h3>
                    <p>Total Words</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"""
                <div class="metric-card">
                    <h3>{analysis['participant_count']}</h3>
                    <p>Participants</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col3:
                st.markdown(f"""
                <div class="metric-card">
                    <h3>{analysis['estimated_duration']}</h3>
                    <p>Est. Duration</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col4:
                st.markdown(f"""
                <div class="metric-card">
                    <h3>{analysis['sentence_count']}</h3>
                    <p>Sentences</p>
                </div>
                """, unsafe_allow_html=True)
            
            if analysis['participants']:
                st.markdown("### 👥 Detected Participants")
                participants_df = []
                for participant, count in analysis['speaking_distribution'].items():
                    total_instances = sum(analysis['speaking_distribution'].values())
                    contribution = (count/total_instances*100) if total_instances > 0 else 0
                    participants_df.append({
                        'Participant': participant,
                        'Speaking Instances': count,
                        'Estimated Contribution': f"{contribution:.1f}%"
                    })
                st.dataframe(participants_df, use_container_width=True)
            
            detected_lang = generator.detect_language(transcript)
            st.info(f"🌐 **Detected Language:** {generator.supported_languages.get(detected_lang, 'Unknown')}")
        else:
            st.info("📝 Enter a transcript in the Text Input tab to see detailed analysis")

    st.divider()
    
    if st.button("🎯 Generate Meeting Minutes", type="primary", use_container_width=True):
        if not transcript:
            st.error("❌ Please provide a meeting transcript")
        else:
            with st.spinner("🤖 Generating professional meeting minutes..."):
                meeting_minutes = generator.generate_meeting_minutes(
                    transcript, selected_language, output_format, meeting_type, custom_instructions
                )
            
            st.session_state.meeting_minutes = meeting_minutes
            st.session_state.transcript_used = transcript
            
            st.markdown("""
            <div class="success-banner">
                <h3>✅ Meeting Minutes Generated Successfully!</h3>
                <p>Your professional meeting minutes are ready for review and export.</p>
            </div>
            """, unsafe_allow_html=True)
            
            result_tab1, result_tab2, result_tab3 = st.tabs(["📄 Minutes Preview", "📊 Summary Stats", "⬇️ Export Options"])
            
            with result_tab1:
                st.markdown("### Generated Meeting Minutes")
                st.text_area(
                    "Meeting Minutes (Read-only preview)",
                    value=meeting_minutes,
                    height=600,
                    disabled=False
                )
                st.code(meeting_minutes, language=None)
            
            with result_tab2:
                st.markdown("### Analysis Summary")
                
                if transcript:
                    analysis = generator.analyze_transcript(transcript)
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("#### 📈 Input Analysis")
                        st.metric("Original Word Count", analysis['word_count'])
                        st.metric("Participants Detected", analysis['participant_count'])
                        st.metric("Estimated Meeting Duration", analysis['estimated_duration'])
                    
                    with col2:
                        st.markdown("#### 📋 Output Summary")
                        minutes_words = len(meeting_minutes.split())
                        compression_ratio = (minutes_words / analysis['word_count']) * 100 if analysis['word_count'] > 0 else 0
                        
                        st.metric("Minutes Word Count", minutes_words)
                        st.metric("Compression Ratio", f"{compression_ratio:.1f}%")
                        st.metric("Format Used", output_format.replace('_', ' ').title())
                
                st.markdown("### ⚙️ Configuration Used")
                config_data = {
                    "Format": output_format.replace('_', ' ').title(),
                    "Meeting Type": meeting_type.replace('_', ' ').title(),
                    "Language": generator.supported_languages[selected_language],
                    "Custom Instructions": custom_instructions if custom_instructions else "None"
                }
                for key, value in config_data.items():
                    st.text(f"{key}: {value}")
            
            with result_tab3:
                st.markdown("### 📥 Export Your Meeting Minutes")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("#### 📄 Text Formats")
                    st.download_button(
                        label="📄 Download as Text (.txt)",
                        data=meeting_minutes,
                        file_name=f"meeting_minutes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                    st.download_button(
                        label="📝 Download as Markdown (.md)",
                        data=meeting_minutes,
                        file_name=f"meeting_minutes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                        mime="text/markdown",
                        use_container_width=True
                    )
                
                with col2:
                    st.markdown("#### 📋 Document Formats")
                    pdf_container = st.container()
                    with pdf_container:
                        if st.button("📋 Generate & Download PDF", key="pdf_export", use_container_width=True):
                            with st.spinner("Creating PDF document..."):
                                pdf_bytes = generator.export_to_pdf(
                                    meeting_minutes,
                                    f"meeting_minutes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                                )
                                if pdf_bytes:
                                    st.success("✅ PDF generated successfully!")
                                    st.download_button(
                                        label="⬇️ Download PDF Now",
                                        data=pdf_bytes,
                                        file_name=f"meeting_minutes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                                        mime="application/pdf",
                                        key="pdf_download_immediate",
                                        use_container_width=True
                                    )
                                else:
                                    st.error("❌ PDF generation failed. Please try again.")
                    docx_container = st.container()
                    with docx_container:
                        if st.button("📄 Generate & Download DOCX", key="docx_export", use_container_width=True):
                            with st.spinner("Creating DOCX document..."):
                                docx_bytes = generator.export_to_docx(
                                    meeting_minutes,
                                    f"meeting_minutes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                                )
                                if docx_bytes:
                                    st.success("✅ DOCX generated successfully!")
                                    st.download_button(
                                        label="⬇️ Download DOCX Now",
                                        data=docx_bytes,
                                        file_name=f"meeting_minutes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="docx_download_immediate",
                                        use_container_width=True
                                    )
                                else:
                                    st.error("❌ DOCX generation failed. Please try again.")
                
                st.divider()
                try:
                    pdf_data = generator.export_to_pdf(meeting_minutes, "meeting_minutes.pdf")
                    if pdf_data:
                        st.download_button(
                            label="📋 Download PDF (Ready)",
                            data=pdf_data,
                            file_name=f"meeting_minutes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf",
                            key="pdf_ready_download",
                            use_container_width=True
                        )
                except Exception as e:
                    st.error(f"PDF error: {str(e)}")
                
                try:
                    docx_data = generator.export_to_docx(meeting_minutes, "meeting_minutes.docx")
                    if docx_data:
                        st.download_button(
                            label="📄 Download DOCX (Ready)",
                            data=docx_data,
                            file_name=f"meeting_minutes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="docx_ready_download",
                            use_container_width=True
                        )
                except Exception as e:
                    st.error(f"DOCX error: {str(e)}")
                
                st.divider()
                json_export = {
                    "meeting_info": {
                        "generated_date": datetime.now().isoformat(),
                        "format": output_format,
                        "meeting_type": meeting_type,
                        "language": selected_language
                    },
                    "transcript": transcript,
                    "meeting_minutes": meeting_minutes,
                    "analysis": generator.analyze_transcript(transcript) if transcript else {}
                }
                st.download_button(
                    label="🔧 Download as JSON (Structured Data)",
                    data=json.dumps(json_export, indent=2),
                    file_name=f"meeting_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json",
                    use_container_width=True
                )
                st.markdown(f"""
                **Configuration Used:**
                - **Format:** {output_format.replace('_', ' ').title()}
                - **Meeting Type:** {meeting_type.replace('_', ' ').title()}
                - **Output Language:** {generator.supported_languages[selected_language]}
                - **Word Count:** {len(meeting_minutes.split())} words
                - **Generated:** {datetime.now().strftime('%Y-%m-%d at %H:%M')}
                """)
                with st.expander("🔧 Download Issues? Try These Solutions"):
                    st.markdown("""
                    **If downloads aren't working:**
                    1. Try the "Quick Download Options" section above
                    2. Check your browser's download settings
                    3. Try a different browser
                    4. Copy the text version from the "Minutes Preview" tab
                    5. Clear browser cache
                    **Alternative Method:**
                    - Copy content from "Minutes Preview" tab
                    - Paste into Google Docs or Microsoft Word
                    - Use the editor's export functions
                    """)

    elif hasattr(st.session_state, 'meeting_minutes'):
        st.info("📋 Previous meeting minutes are available from your session")
        
        with st.expander("View Previous Results"):
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.download_button(
                    label="📄 Download as Text",
                    data=st.session_state.meeting_minutes,
                    file_name=f"meeting_minutes_previous.txt",
                    mime="text/plain"
                )
            
            with col2:
                st.download_button(
                    label="📝 Download as Markdown",
                    data=st.session_state.meeting_minutes,
                    file_name=f"meeting_minutes_previous.md",
                    mime="text/markdown"
                )
            
            with col3:
                try:
                    pdf_data = generator.export_to_pdf(st.session_state.meeting_minutes, "meeting_minutes_previous.pdf")
                    if pdf_data:
                        st.download_button(
                            label="📋 Download as PDF",
                            data=pdf_data,
                            file_name=f"meeting_minutes_previous.pdf",
                            mime="application/pdf"
                        )
                except Exception as e:
                    st.error(f"PDF error: {str(e)}")
            
            with col4:
                try:
                    docx_data = generator.export_to_docx(st.session_state.meeting_minutes, "meeting_minutes_previous.docx")
                    if docx_data:
                        st.download_button(
                            label="📄 Download as DOCX",
                            data=docx_data,
                            file_name=f"meeting_minutes_previous.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"DOCX error: {str(e)}")
            
            if st.button("👁️ Preview Previous Minutes"):
                st.text_area("Previous Meeting Minutes", st.session_state.meeting_minutes, height=400)

    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; padding: 2rem; background: #f8f9fa; border-radius: 10px; margin-top: 2rem;'>
        <h4>🤖 AI-Powered Meeting Documentation</h4>
        <p><strong>Features:</strong> Advanced AI Processing | Multi-Language Support | Professional Export Options<br><strong>Powered by:</strong> Mistral-Nemo AI | Google Speech Recognition | Advanced Document Processing</p>
        <p style='color: #666; font-size: 0.9em;'>Generated with ❤️ for professional meeting documentation</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()