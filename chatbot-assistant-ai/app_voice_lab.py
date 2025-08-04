# Audio & Text Processing Tool
import streamlit as st
import whisper
import tempfile
import os
from TTS.api import TTS
from pydub import AudioSegment
from io import BytesIO
from docx import Document
import time
import numpy as np
from scipy.io.wavfile import write as wav_write
import textwrap

# Custom CSS
st.markdown("""
    <style>
    .main-header { 
        text-align: center; 
        padding: 1.5rem; 
        background: linear-gradient(90deg, #4a90e2, #9013fe); 
        color: white; 
        border-radius: 10px; 
        margin-bottom: 1.5rem; 
    }
    .tab-content { 
        padding: 1rem; 
        border-radius: 8px; 
        background: #f9f9f9; 
    }
    .success-message { 
        background-color: #d4edda; 
        color: #155724; 
        padding: 1rem; 
        border-radius: 5px; 
        margin-top: 1rem; 
    }
    .error-message { 
        background-color: #f8d7da; 
        color: #721c24; 
        padding: 1rem; 
        border-radius: 5px; 
        margin-top: 1rem; 
    }
    </style>
""", unsafe_allow_html=True)

# Initialize session state
if 'transcription_result' not in st.session_state:
    st.session_state.transcription_result = ""
if 'tts_audio' not in st.session_state:
    st.session_state.tts_audio = None

# Load models
@st.cache_resource
def load_whisper_model():
    return whisper.load_model("base")

@st.cache_resource
def load_tts_model():
    return TTS(model_name="tts_models/multilingual/multi-dataset/your_tts", progress_bar=False, gpu=False)

whisper_model = load_whisper_model()
tts_model = load_tts_model()

# Supported languages
transcription_langs = {
    "en": "English", "fr": "French", "es": "Spanish", "de": "German",
    "it": "Italian", "pt": "Portuguese", "nl": "Dutch", "ru": "Russian"
}
tts_langs = ["en"]

# Transcription function
def transcribe_audio(file_path, lang_code):
    try:
        result = whisper_model.transcribe(file_path, language=lang_code)
        return result["text"]
    except Exception as e:
        st.error(f"Transcription error: {str(e)}")
        return ""

# ✅ FIXED: Smart text chunking — preserves full meaning
def chunk_text(text, max_chars=600):
    text = " ".join(text.split())  # Normalize whitespace
    return [
        chunk.strip() 
        for chunk in textwrap.wrap(text, width=max_chars, break_long_words=True, replace_whitespace=False)
        if chunk.strip()
    ]

# Extract audio and sample rate from TTS output
def extract_audio_and_sr(audio_out):
    sample_rate = getattr(tts_model, "output_sample_rate", 22050)
    if isinstance(audio_out, dict):
        return audio_out.get("wav", list(audio_out.values())[0]), sample_rate
    elif isinstance(audio_out, tuple) and len(audio_out) == 2:
        return audio_out[0], audio_out[1]
    elif isinstance(audio_out, (list, np.ndarray)):
        return audio_out, sample_rate
    else:
        raise ValueError(f"Unsupported TTS output: {type(audio_out)}")

# ✅ TTS synthesis: Now speaks ALL text using pydub for safe concatenation
def synthesize_text(text, speaker, lang, chunk_size, ref_audio=None):
    try:
        if not text.strip():
            raise ValueError("Empty text")

        chunks = chunk_text(text, int(chunk_size))
        st.write(f"🔤 Split into {len(chunks)} chunks. Synthesizing each...")

        # Use pydub to concatenate audio properly
        combined_audio = AudioSegment.empty()
        silence = AudioSegment.silent(duration=300)  # 300ms silence between chunks

        sample_rate = 22050
        for i, chunk in enumerate(chunks):
            st.write(f"🔊 Chunk {i+1}/{len(chunks)}: {len(chunk)} chars")
            try:
                audio_out = tts_model.tts(
                    text=chunk,
                    speaker=speaker,
                    language=lang,
                    speaker_wav=ref_audio
                )
                audio, sr = extract_audio_and_sr(audio_out)
                sample_rate = sr

                # Convert to numpy array
                if isinstance(audio, list):
                    audio = np.array(audio, dtype=np.float32)
                audio = audio.astype(np.float32)

                # Normalize
                if np.max(np.abs(audio)) > 1e-5:
                    audio = audio / np.max(np.abs(audio))

                # Convert to 16-bit PCM
                audio_int16 = (audio * 32767).astype(np.int16)

                # Save to BytesIO and load with pydub
                buf = BytesIO()
                wav_write(buf, rate=sample_rate, data=audio_int16)
                buf.seek(0)
                chunk_audio = AudioSegment.from_file(buf, format="wav")

                # Append to final audio
                if len(combined_audio) > 0:
                    combined_audio += silence
                combined_audio += chunk_audio

            except Exception as e:
                st.warning(f"⚠️ Failed on chunk {i+1}: {str(e)}")
                continue  # Skip only this one

        if len(combined_audio) == 0:
            st.error("❌ No audio was generated.")
            return None

        # Export final audio to MP3 via BytesIO
        mp3_buffer = BytesIO()
        combined_audio.export(mp3_buffer, format="mp3", bitrate="192k")
        mp3_buffer.seek(0)
        return mp3_buffer.read()

    except Exception as e:
        st.error(f"🛑 TTS synthesis error: {str(e)}")
        return None

# Export to DOCX
def export_to_docx(text, filename):
    doc = Document()
    doc.add_heading("Transcription", 0)
    doc.add_paragraph(text)
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()

# ===========================
# Main App UI
# ===========================

st.markdown('<div class="main-header"><h1>🎙️ Audio & Text Processing Tool</h1><p>Transcribe MP3 to Text & Convert Text to Speech</p></div>', unsafe_allow_html=True)

tab1, tab2 = st.tabs(["🎙️ Audio Transcription", "🔊 Text to Speech"])

# --------------------------
# TAB 1: Audio Transcription
# --------------------------
with tab1:
    st.markdown('<div class="tab-content">', unsafe_allow_html=True)
    lang_code = st.selectbox(
        "Select Language for Transcription",
        options=list(transcription_langs.keys()),
        format_func=lambda x: transcription_langs[x],
        key="transcribe_lang"
    )
    audio_file = st.file_uploader("Upload your MP3 audio file", type=["mp3"])

    if audio_file:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
            tmp_file.write(audio_file.read())
            temp_path = tmp_file.name
        st.audio(temp_path)

        if st.button("Transcribe Audio"):
            with st.spinner("Transcribing..."):
                time.sleep(1)
                result = transcribe_audio(temp_path, lang_code)
                st.session_state.transcription_result = result
                if result:
                    st.markdown('<div class="success-message">✅ Transcription completed!</div>', unsafe_allow_html=True)
                else:
                    st.markdown('<div class="error-message">❌ Failed.</div>', unsafe_allow_html=True)
            try:
                os.unlink(temp_path)
            except:
                pass

    if st.session_state.transcription_result:
        st.subheader("📝 Transcribed Text")
        st.text_area("Result", value=st.session_state.transcription_result, height=300)
        col1, col2 = st.columns(2)
        with col1:
            st.download_button("📥 TXT", data=st.session_state.transcription_result, file_name="transcription.txt", mime="text/plain")
        with col2:
            docx_data = export_to_docx(st.session_state.transcription_result, "transcription.docx")
            st.download_button("📄 DOCX", data=docx_data, file_name="transcription.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.markdown('</div>', unsafe_allow_html=True)

# --------------------------
# TAB 2: Text to Speech
# --------------------------
with tab2:
    st.markdown('<div class="tab-content">', unsafe_allow_html=True)
    text = st.text_area("📝 Enter your text (ALL will be spoken)", height=250, key="tts_input")
    speaker = st.selectbox("🗣️ Voice", options=tts_model.speakers if hasattr(tts_model, 'speakers') else ["default"])
    lang = st.selectbox("🌐 Language", options=tts_langs, format_func=lambda x: "English", index=0)
    chunk_size = st.number_input("🔤 Max chunk length", min_value=100, max_value=1000, value=600, key="chunk_size")

    uploaded_audio = st.file_uploader("🎚️ Voice clone reference (WAV/MP3)", type=["wav", "mp3"], key="ref_audio")

    ref_audio_path = None
    if uploaded_audio:
        with tempfile.NamedTemporaryFile(delete=False, suffix="." + uploaded_audio.name.split(".")[-1]) as tmp:
            tmp.write(uploaded_audio.read())
            ref_audio_path = tmp.name

    if st.button("🔊 Convert to Speech"):
        if not text.strip():
            st.markdown('<div class="error-message">❌ Enter text!</div>', unsafe_allow_html=True)
        else:
            with st.spinner("🎙️ Generating full audio..."):
                time.sleep(1)
                mp3_audio = synthesize_text(text, speaker, lang, chunk_size, ref_audio_path)
                if ref_audio_path:
                    try:
                        os.unlink(ref_audio_path)
                    except:
                        pass
                if mp3_audio:
                    st.session_state.tts_audio = mp3_audio
                    st.markdown('<div class="success-message">✅ Full text spoken successfully!</div>', unsafe_allow_html=True)
                    st.audio(mp3_audio, format="audio/mp3")
                    st.download_button("💾 Save as MP3", data=mp3_audio, file_name="full_speech.mp3", mime="audio/mp3")
                else:
                    st.markdown('<div class="error-message">❌ Synthesis failed.</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

# Footer
st.markdown(
    '<div style="text-align: center; padding: 1rem; color: #666;">'
    'Powered by <strong>Whisper</strong> & <strong>Coqui TTS</strong> | © 2025'
    '</div>',
    unsafe_allow_html=True
)