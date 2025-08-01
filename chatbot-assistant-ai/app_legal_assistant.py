# AI-Powered Legal Assistant mistral ollama
import streamlit as st
import requests
from docx import Document
from fpdf import FPDF
import io

# ----------------------------
# Configuration
# ----------------------------
OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL_NAME = "mistral-nemo:latest"

# ----------------------------
# Multilingual Legal Templates
# ----------------------------
LEGAL_TEMPLATES = {
    "rental agreement": {
        "en": "Generate a rental agreement between {party1} (tenant) and {party2} (landlord) for {duration} months.",
        "fr": "Rédigez un contrat de location entre {party1} (locataire) et {party2} (bailleur) pour une durée de {duration} mois.",
        "es": "Elabora un contrato de arrendamiento entre {party1} (arrendatario) y {party2} (arrendador) por {duration} meses."
    },
    "employment contract": {
        "en": "Generate an employment contract between {party1} (employee) and {party2} (employer) with a salary of {salary} per year.",
        "fr": "Rédigez un contrat de travail entre {party1} (employé) et {party2} (employeur) avec un salaire de {salary} par an.",
        "es": "Elabora un contrato laboral entre {party1} (empleado) y {party2} (empleador) con un salario de {salary} anual."
    },
    "business partnership agreement": {
        "en": "Draft a business partnership agreement between {party1} and {party2}, defining responsibilities and profit-sharing terms.",
        "fr": "Rédigez un accord de partenariat commercial entre {party1} et {party2}, définissant les responsabilités et les modalités de partage des bénéfices.",
        "es": "Redacta un acuerdo de sociedad comercial entre {party1} y {party2}, definiendo responsabilidades y términos de reparto de beneficios."
    },
    "nda": {
        "en": "Generate a non-disclosure agreement (NDA) between {party1} and {party2} to protect confidential business information.",
        "fr": "Rédigez un accord de confidentialité (NDA) entre {party1} et {party2} pour protéger les informations commerciales confidentielles.",
        "es": "Elabora un acuerdo de confidencialidad (NDA) entre {party1} y {party2} para proteger información comercial confidencial."
    },
    "freelance contract": {
        "en": """
        Generate a professional freelance contract between {party1} (freelancer) and {party2} (client) for the project: "{scope}".
        Duration: {duration} months. Payment: {salary}.
        Include clauses for deliverables, revisions, payment terms, ownership of work, and termination.
        """,
        "fr": """
        Rédigez un contrat de freelance professionnel entre {party1} (freelance) et {party2} (client) pour le projet : « {scope} ».
        Durée : {duration} mois. Rémunération : {salary}.
        Incluez des clauses sur les livrables, les révisions, les conditions de paiement, la propriété du travail et la résiliation.
        """,
        "es": """
        Elabora un contrato de trabajo freelance entre {party1} (freelance) y {party2} (cliente) para el proyecto: "{scope}".
        Duración: {duration} meses. Pago: {salary}.
        Incluye cláusulas sobre entregables, revisiones, términos de pago, propiedad del trabajo y terminación.
        """
    }
}

LANGUAGE_NAMES = {
    "en": "English",
    "fr": "Français",
    "es": "Español"
}

# ----------------------------
# Helper: Call Ollama API
# ----------------------------
def generate_legal_document(
    doc_type: str,
    party1: str,
    party2: str,
    duration: str = "",
    salary: str = "",
    lang: str = "en",
    scope: str = ""
) -> str:
    template = LEGAL_TEMPLATES.get(doc_type.lower(), {}).get(lang, "")
    if not template:
        return {"en": "Invalid document type.", "fr": "Type de document invalide.", "es": "Tipo de documento inválido."}[lang]

    try:
        prompt = template.format(party1=party1, party2=party2, duration=duration, salary=salary, scope=scope)
    except KeyError as e:
        missing = str(e)
        return {
            "en": f"Missing required field: {missing}",
            "fr": f"Champ requis manquant : {missing}",
            "es": f"Falta el campo requerido: {missing}"
        }[lang]

    payload = {
        "model": MODEL_NAME,
        "prompt": prompt,
        "stream": False
    }

    try:
        response = requests.post(OLLAMA_URL, json=payload)
        if response.status_code == 200:
            return response.json().get("response", "No content generated.")
        else:
            return f"Error {response.status_code}: {response.text}"
    except requests.exceptions.RequestException as e:
        return f"Request failed: {str(e)}"

# ----------------------------
# Export Functions
# ----------------------------
def create_docx(text: str) -> io.BytesIO:
    doc = Document()
    doc.add_heading('Generated Legal Document', 0)
    for line in text.split('\n'):
        line = line.strip()
        if line:
            doc.add_paragraph(line)
    byte_stream = io.BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)
    return byte_stream

def create_pdf(text: str) -> io.BytesIO:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.set_auto_page_break(auto=True, margin=15)
    # Add title
    pdf.set_font("Arial", style="B", size=16)
    pdf.cell(0, 12, "Generated Legal Document", ln=True, align="C")
    pdf.ln(8)
    pdf.set_font("Arial", size=12)
    for line in text.split('\n'):
        pdf.multi_cell(0, 10, line)
    pdf_buffer = io.BytesIO()
    pdf_data = pdf.output(dest='S').encode('latin-1')
    pdf_buffer.write(pdf_data)
    pdf_buffer.seek(0)
    return pdf_buffer

# ----------------------------
# Streamlit UI
# ----------------------------
st.set_page_config(
    page_title="AI Legal Assistant",
    page_icon="⚖️",
    layout="centered"
)

st.title("⚖️ AI-Powered Legal Assistant")
st.markdown("""
**Generate professional legal contracts instantly.**  
Powered by `mistral-nemo:latest` via Ollama. Supports multiple languages.
""")

# Language selector
language = st.selectbox(
    "🌐 Select Language",
    options=list(LANGUAGE_NAMES.keys()),
    format_func=lambda x: LANGUAGE_NAMES[x],
    index=0
)

# Document type
doc_type_label = {
    "en": "Document Type",
    "fr": "Type de document",
    "es": "Tipo de documento"
}[language]
doc_type = st.selectbox(
    doc_type_label,
    options=list(LEGAL_TEMPLATES.keys()),
    format_func=lambda x: x.replace("-", " ").title()
).lower()

# Parties
party1_label = {"en": "Party 1 Name", "fr": "Nom de la partie 1", "es": "Nombre de la parte 1"}[language]
party2_label = {"en": "Party 2 Name", "fr": "Nom de la partie 2", "es": "Nombre de la parte 2"}[language]
party1 = st.text_input(party1_label, placeholder="e.g., John Doe or Company Inc.")
party2 = st.text_input(party2_label, placeholder="e.g., Jane Smith or Client LLC")

# Optional fields based on document type
duration = ""
salary = ""
scope = ""

if doc_type == "rental agreement":
    duration_label = {"en": "Duration (months)", "fr": "Durée (mois)", "es": "Duración (meses)"}[language]
    duration = st.text_input(duration_label, placeholder="e.g., 12")

elif doc_type == "employment contract":
    salary_label = {"en": "Annual Salary", "fr": "Salaire annuel", "es": "Salario anual"}[language]
    salary = st.text_input(salary_label, placeholder="e.g., $75,000")

elif doc_type == "freelance contract":
    duration_label = {"en": "Project Duration (months)", "fr": "Durée du projet (mois)", "es": "Duración del proyecto (meses)"}[language]
    duration = st.text_input(duration_label, placeholder="e.g., 3")

    salary_label = {"en": "Total Fee or Rate", "fr": "Honoraires totaux ou tarif", "es": "Honorarios totales o tarifa"}[language]
    salary = st.text_input(salary_label, placeholder="e.g., $5,000 or $50/hour")

    scope_label = {"en": "Project Scope / Deliverables", "fr": "Périmètre du projet / Livrables", "es": "Alcance del proyecto / Entregables"}[language]
    scope = st.text_area(scope_label, placeholder="e.g., Design a responsive website with SEO and 5 pages.")

# Generate button
if st.button("📄 Generate Contract"):
    if not party1.strip() or not party2.strip():
        st.error({"en": "Please enter both party names.", "fr": "Veuillez entrer les noms des deux parties.", "es": "Por favor, ingrese los nombres de ambas partes."}[language])
    else:
        with st.spinner({"en": "Generating contract...", "fr": "Génération du contrat...", "es": "Generando contrato..."}[language]):
            generated_text = generate_legal_document(doc_type, party1, party2, duration, salary, language, scope=scope)
            st.session_state.generated_text = generated_text
            st.subheader("📝 Generated Document")
            st.text_area("", value=generated_text, height=400)

# ----------------------------
# Download Options
# ----------------------------
if "generated_text" in st.session_state and st.session_state.generated_text:
    st.markdown("### 📥 Download Document")

    col1, col2 = st.columns(2)

    # Download as Word (.docx)
    with col1:
        docx_file = create_docx(st.session_state.generated_text)
        st.download_button(
            label="⬇️ Download as Word (.docx)",
            data=docx_file,
            file_name="legal_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Download as PDF
    with col2:
        pdf_file = create_pdf(st.session_state.generated_text)
        st.download_button(
            label="🖨️ Download as PDF",
            data=pdf_file,
            file_name="legal_document.pdf",
            mime="application/pdf"
        )

# ----------------------------
# Footer / License Note
# ----------------------------
st.markdown("---")
st.markdown("""
### 📄 License & Disclaimer

This app is for **informational and educational purposes only**.  
Generated content is not legal advice. Always consult a licensed attorney.

**Freemium Use**: You may use, modify, and share this code freely (MIT-style permissive use).  
No warranty. Use at your own risk.

Powered by [Ollama](https://ollama.com) + `mistral-nemo:latest`.
""")